"""Content extraction for each knowledge source type — spec §16 ingestion
pipeline's first stage. docx/csv extraction is a real, working narrow slice
(python-docx paragraph/table text, naive CSV-to-text); OCR'd/scanned PDFs and
complex table layouts are out of scope for this pass.
"""

from __future__ import annotations

import csv
import io
from urllib.parse import urlparse

import httpx
from bs4 import BeautifulSoup
from docx import Document
from fastapi import HTTPException, status
from pypdf import PdfReader

from app.security import assert_public_host


async def fetch_website_text(url: str, *, max_redirects: int = 3, max_bytes: int = 5_000_000) -> tuple[str, str]:
    """Returns (title, text). Redirects are followed manually, one hop at a
    time, re-validating the target host each time — httpx's built-in
    follow_redirects would only re-check DNS at the very end, not each hop."""
    current_url = url
    async with httpx.AsyncClient(timeout=15.0, follow_redirects=False) as client:
        for _ in range(max_redirects + 1):
            parsed = urlparse(current_url)
            if parsed.scheme not in ("http", "https"):
                raise HTTPException(status.HTTP_400_BAD_REQUEST, "Only http/https URLs are supported")
            if not parsed.hostname:
                raise HTTPException(status.HTTP_400_BAD_REQUEST, "Invalid URL")
            assert_public_host(parsed.hostname, current_url)

            response = await client.get(current_url, headers={"User-Agent": "JKRAICallingBot/1.0"})
            if response.status_code in (301, 302, 303, 307, 308):
                location = response.headers.get("location")
                if not location:
                    raise HTTPException(status.HTTP_502_BAD_GATEWAY, "Redirect with no Location header")
                current_url = httpx.URL(current_url).join(location).human_repr()
                continue

            response.raise_for_status()
            content = response.content[:max_bytes]
            soup = BeautifulSoup(content, "html.parser")
            for tag in soup(["script", "style", "nav", "footer", "header"]):
                tag.decompose()
            title = soup.title.string.strip() if soup.title and soup.title.string else current_url
            text = " ".join(soup.get_text(separator=" ").split())
            return title, text

    raise HTTPException(status.HTTP_400_BAD_REQUEST, "Too many redirects")


def extract_text_from_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(p.strip() for p in pages if p.strip())


def extract_text_from_docx(data: bytes) -> str:
    document = Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def extract_text_from_csv(data: bytes) -> str:
    text = data.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    rows = [", ".join(cell.strip() for cell in row) for row in reader if any(c.strip() for c in row)]
    return "\n".join(rows)


def chunk_text(text: str, *, max_chars: int = 800, overlap_chars: int = 100) -> list[str]:
    """Paragraph-aware chunking: merges paragraphs until the size limit, then
    starts a new chunk carrying a small overlap forward so a fact split
    across a chunk boundary isn't lost to either side."""
    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    if not paragraphs:
        return []

    chunks: list[str] = []
    current = ""
    for para in paragraphs:
        candidate = f"{current}\n{para}".strip() if current else para
        if len(candidate) <= max_chars:
            current = candidate
            continue
        if current:
            chunks.append(current)
            tail = current[-overlap_chars:] if overlap_chars else ""
            current = f"{tail}\n{para}".strip()
        else:
            # A single paragraph longer than max_chars — hard-split it.
            for i in range(0, len(para), max_chars):
                chunks.append(para[i : i + max_chars])
            current = ""
    if current:
        chunks.append(current)
    return chunks
